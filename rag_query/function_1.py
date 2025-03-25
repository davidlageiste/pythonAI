import azure.functions as func
import logging
from langchain.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader, JSONLoader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
import faiss
from langchain.schema import Document
from openai import OpenAI
from docx import Document as DocxDocument
import os
import json
import tempfile
import pickle
from azure.storage.blob import BlobServiceClient

os.environ['FAISS_NO_GPU'] = '1'
# Configuration du logger
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

class KnowledgeBase:
    def __init__(self, text_data):
        self.embeddings = OpenAIEmbeddings(model="text-embedding-ada-002",api_key= os.environ["OPENAI_API_KEY"])
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size =200, chunk_overlap =50,separators=["\n"])
        self.documents=self.convert_texts_to_documents(self.text_split(text_data)) if text_data else []
        self.retriever = None
        
    def convert_texts_to_documents(self, texts):
        return [Document(page_content=text) for text in texts]
        
    def text_split(self , extracted_data):
        return self.text_splitter.split_text(extracted_data)

    def build_retriever(self, connection_string, container_name):
        store = FAISS.from_documents(self.documents, self.embeddings)
        self._save_to_blob(store, connection_string, container_name)
        self.retriever = store.as_retriever()

    def _save_to_blob(self , store, connection_string, container_name):
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        with tempfile.NamedTemporaryFile(delete=False) as temp_index:
            faiss.write_index(store.index, temp_index.name)
            with open(temp_index.name, 'rb') as data:
                container_client.upload_blob(name="index.faiss", data=data, overwrite=True)
            os.unlink(temp_index.name)
        with tempfile.NamedTemporaryFile(delete=False) as temp_docstore:
            pickle.dump(store.docstore, temp_docstore)
            temp_docstore.flush()
            with open(temp_docstore.name, 'rb') as data:
                container_client.upload_blob(name="docstore.pkl", data=data, overwrite=True)
            os.unlink(temp_docstore.name)



class VirtualAssistant:
    def __init__(self, retriever, llm_model="gpt-3.5-turbo"):
        self.retriever = retriever
        self.client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    def answer_query(self, query, context):
        custom_prompt_template = (
            f"Vous êtes une assistante virtuelle spécialisée pour un cabinet de radiologie médicale.\n"
            f"Voici des informations pertinentes extraites de notre base de connaissances :\n{context}\n\n"
            f"Question du patient : {query}\n\n"
            f"Répondez uniquement à la question posée, de manière claire et concise."
        )

        completion = self.client.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[
                {"role": "system", "content": custom_prompt_template},
                {"role": "user", "content": query}
            ]
        )
        return completion.choices[0].message.content

class RAG_Azure:
    def __init__(self , llm_model="gpt-3.5-turbo"):
        self.connection_string = os.environ["AZURE_STORAGE_CONNECTION_STRING"]
        self.container_name = os.environ["AZURE_STORAGE_CONTAINER_NAME"]
        self.knowledge_base = KnowledgeBase('')
        logger.info(f"knowledge base from empty string is well done")
        try:
            self.knowledge_base.retriever = self._load_from_blob().as_retriever()
            logger.info(f"Could  load from blob storage: {str(e)}. No Building new index.")
        except Exception as e:
            logger.info(f"Could not load from blob storage: {str(e)}. Building new index.")
            self.knowledge_base = KnowledgeBase(self.load_files_contents('data'))
            self.knowledge_base.build_retriever(self.connection_string, self.container_name)

        self.assistant = VirtualAssistant(
            retriever=self.knowledge_base.retriever,
            llm_model=llm_model
        )

    def _load_from_blob(self):
        try:
            logger.info("Démarrage du chargement depuis Azure Blob Storage.")
    
            # Connexion au Blob Storage
            logger.info("Connexion à Azure Blob Storage.")
            blob_service_client = BlobServiceClient.from_connection_string(self.connection_string)
            container_client = blob_service_client.get_container_client(self.container_name)
    
            # Récupération des blobs
            logger.info("Récupération des blobs 'index.faiss' et 'docstore.pkl'.")
            index_blob_client = container_client.get_blob_client("index.faiss")
            docstore_blob_client = container_client.get_blob_client("docstore.pkl")
    
            # Vérification de l'existence des blobs
            if not index_blob_client.exists():
                logger.error("Le blob 'index.faiss' est introuvable.")
                raise FileNotFoundError("Le blob 'index.faiss' est introuvable.")
            if not docstore_blob_client.exists():
                logger.error("Le blob 'docstore.pkl' est introuvable.")
                raise FileNotFoundError("Le blob 'docstore.pkl' est introuvable.")
    
            # Téléchargement et écriture des fichiers temporaires
            with tempfile.NamedTemporaryFile(delete=False) as temp_index:
                logger.info("Téléchargement et écriture du fichier temporaire pour 'index.faiss'.")
                temp_index.write(index_blob_client.download_blob().readall())
                temp_index.close()
    
            with tempfile.NamedTemporaryFile(delete=False) as temp_docstore:
                logger.info("Téléchargement et écriture du fichier temporaire pour 'docstore.pkl'.")
                temp_docstore.write(docstore_blob_client.download_blob().readall())
                temp_docstore.close()
    
            # Chargement de l'index FAISS
            logger.info("Chargement de l'index FAISS depuis le fichier temporaire.")
            index = faiss.read_index(temp_index.name)
    
            # Chargement du docstore
            logger.info("Chargement du docstore depuis le fichier temporaire.")
            with open(temp_docstore.name, 'rb') as f:
                docstore = pickle.load(f)
    
            # Nettoyage des fichiers temporaires
            logger.info("Suppression des fichiers temporaires.")
            os.unlink(temp_index.name)
            os.unlink(temp_docstore.name)
    
            logger.info("Chargement terminé avec succès.")
            return FAISS(embedding_function=self.knowledge_base.embeddings.embed_query, index=index, docstore=docstore ,index_to_docstore_id = {})
    
        except Exception as e:
            logger.exception(f"Erreur lors du chargement des blobs : {e}")
            raise RuntimeError(f"Erreur lors du chargement des blobs : {e}")

    
    def load_files_contents(self , data_path):
            content_parts = []  
            file_handlers = {
                ".pdf": lambda path: "\n".join(doc.page_content for doc in PyPDFLoader(path).load()),
                ".txt": lambda path: open(path, 'r', encoding='utf-8').read(),
                ".json": lambda path: json.dumps(json.load(open(path, 'r', encoding='utf-8')), indent=2),
                ".docx": lambda path: "\n".join(paragraph.text for paragraph in DocxDocument(path).paragraphs),
            }
            for filename in os.listdir(data_path):
                filepath = os.path.join(data_path, filename)
                file_ext = os.path.splitext(filename)[1]  
                if file_ext in file_handlers:
                    try:
                        content_parts.append(file_handlers[file_ext](filepath))
                    except Exception as e:
                        logger.error(f"Error loading {filepath}: {e}")
            return "\n".join(content_parts)  

    def process_query(self, query):
        retrieved_docs = self.knowledge_base.retriever.get_relevant_documents(query , k=4)
        if not retrieved_docs:
            return "Aucun document pertinent trouvé."
        context ="\n".join(doc.page_content for doc in retrieved_docs) 
        return self.assistant.answer_query(query, context)
