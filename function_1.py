from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
from sentence_transformers import SentenceTransformer
import re
import dateparser
import requests
import json
import faiss
import os
if not hasattr(os, 'add_dll_directory'):
    os.add_dll_directory = lambda x: None

import unicodedata
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.chat_models import ChatOpenAI
from langchain.schema import Document
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader , JSONLoader
from langchain.embeddings import HuggingFaceEmbeddings
from docx import Document as DocxDocument
from datetime import datetime
import logging


logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
ch = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
ch.setFormatter(formatter)
logger.addHandler(ch)


class AnalyseurConversation:
    def __init__(self, model_name='paraphrase-MiniLM-L6-v2'):
        exemples = {
            "oui": ["Oui, bien sûr.", "yes", "je suis ok", "D'accord, je suis partant.", "Oui, je le fais volontiers.", "C'est une excellente idée.", "Oui, sans hésiter.", "Oui, je confirme.", "Bien sûr, je suis d'accord.", "Oui, c'est tout à fait correct.", "Je suis pour.", "Oui, je suis avec vous.", "Sans problème, c'est oui.", "Oui, je suis partiellement d'accord.", "Oui, pourquoi pas.", "D'accord, je le ferai.", "Oui, c'est une bonne solution.", "Bien sûr, je le fais sans hésiter.", "Je consens", "je veux", "Absolument", "Avec plaisir", "Oui, je suis d'accord", "Ça me va", "Bien entendu", "Pas de souci", "Ok, c'est bon", "C'est bon pour moi", "Je suis favorable", "C'est d'accord", "Tout à fait", "J'accepte"],
            "non": ["Non, merci.", "no", "Je ne suis pas d'accord.", "Non, ce n'est pas pour moi.", "Je préfère ne pas.", "Non, ce n'est pas possible.", "Non, je ne veux pas.", "Je ne suis pas intéressé.", "Non, je ne le ferai pas.", "Ce n'est pas ce que je veux.", "Non, je ne crois pas.", "Non, je ne pense pas.", "Non, c'est non.", "Je m'abstiens.", "Non, pas question.", "Ce n'est pas acceptable pour moi.", "Non, je refuse catégoriquement.", "Je ne consens pas", "je ne veux pas", "Pas du tout", "C'est non", "Je ne peux pas", "Ce n'est pas possible", "Je ne suis pas d'accord avec ça", "C'est hors de question", "Je refuse", "Non, ça ne m'intéresse pas", "Je décline"],
            "indéterminé": ["Je ne suis pas sûr.", "Je ne sais pas.", "Peut-être, je ne sais pas vraiment.", "Je ne suis pas convaincu.", "C'est compliqué.", "Je doute.", "Ça m'embête.", "Je ne suis pas certain.", "C'est flou.", "Je crois que non.", "Je ne suis pas certain de ma réponse.", "Je ne sais pas quoi répondre.", "Je suis hésitant.", "C'est ambigu.", "Je ne suis pas clair sur ma réponse.", "Je n'ai pas d'avis.", "Je ne sais pas trop.", "Je suis indécis.", "C'est un peu flou pour moi.", "Je suis partagé.", "J'ai des doutes.", "Je n'ai pas de réponse précise.", "Je ne peux pas me prononcer.", "C'est incertain.", "Je suis perplexe.", "Je n'ai pas de certitude.", "C'est difficile à dire.", "Je n'ai pas d'opinion claire."]
        }

        self.model = SentenceTransformer(model_name)
        self.init_faiss_index(exemples)
        self.pattern_quitter = r"""\b(quitt(?:e|er|é|ant)?|part(?:i|ir|ait|ie|is|ons)?|arrêt(?:er|e|ons)?|fini(?:r|e|s)?|stop|au revoir|termin(?:er|é|e)?|m'en vais|ferm(?:er|é|ée|ons)?|bientôt|clôtur(?:er|e|ons)?|fin(?:ir|ie)?|c'est tout|ça y est|je file|je m'en vais|je dois y aller|je me tire|je me casse|je bounce|j'y vais|bon, j'y vais|go)\b"""

    def init_faiss_index(self, exemples):
        phrases = []
        categories = []

        for category, examples in exemples.items():
            for phrase in examples:
                phrases.append(phrase.lower())
                categories.append(category)
        embeddings = self.model.encode(phrases)
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(embeddings)
        self.categories = categories

    def quitter_conversation(self, question):
        return bool(re.search(self.pattern_quitter, question.lower()))


    def classer_reponse(self, reponse):
        question_embedding = self.model.encode([reponse.lower()])
        distances, indices = self.index.search(question_embedding, 1)
        category = self.categories[indices[0][0]]
        if category == "oui":
            return 'oui'
        elif category == "non":
            return 'non'
        else:
            return "indéterminé"

    def recueil_consentement(self, reponse):
        classification = self.classer_reponse(reponse)
        if classification == "oui":
            return {
                "status": "Consentement valide",
                "réponse": "Nous vous remercions pour votre consentement. Maintenant , nous pouvons continuer."
            }
        elif classification == "non":
            return {
                "status": "Consentement non valide",
                "réponse": "Nous avons bien pris en compte votre refus de consentement. Dans ce cas, nous vous invitons à contacter le cabinet dentaire directement ou à joindre votre dentiste en appelant le numéro du cabinet."
            }
        elif classification == "indéterminé":
            return {
                "status": "Consentement indéterminé",
                "réponse": "Votre réponse est ambigüe, nécessitant une clarification pour valider ou invalider le consentement."
            }
        else:
            return {
                "status": "Réponse non catégorisée",
                "réponse": "La réponse n'a pas pu être catégorisée. Veuillez répondre oui ou non."
            }

class VirtualAssistant:
    def __init__(self, retriever,general=True, llm_model="gpt-3.5-turbo"):
        self.retriever = retriever
        self.client = OpenAI(api_key = os.environ["OPENAI_API_KEY"])
        self.general=general

    def answer_query(self, query,context ):
        custom_prompt_template_general= (
            f"Vous êtes une assistante virtuelle spécialisée pour un cabinet de radiologie médicale.Ton role consiste à répondre aux questions des patients \n"
            f"Voici des informations pertinentes extraites de notre base de connaissances :\n{context}\n\n"
            f"Question du patient : {query}\n\n"
            f"Répondez uniquement à la question posée, de manière claire et concise et avec le langage naturel et simple, en utilisant uniquement les informations fournies ci-dessus. "
            f"N'ajoutez pas d'autres informations ou détails, et signalez si une donnée spécifique est absente."
        )
        custom_prompt_template_personnal = (
            f"Vous êtes une assistante virtuelle dédiée à un cabinet de radiologie médicale, spécialisée dans l'accompagnement des patients.\n"
            f"Votre rôle est d'apporter des réponses personnalisées aux questions des patients, en utilisant les informations suivantes extraites de notre base de données médicales :\n{context}\n\n"
            f"Voici la question du patient : {query}\n\n"
            f"Répondez de manière claire, concise et avec un langage simple et naturel, adapté à un patient. Utilisez uniquement les informations fournies ci-dessus. "
            f"Si une donnée spécifique est manquante ou insuffisante pour répondre à la question, signalez-le clairement. "
            f"Assurez-vous que votre réponse reste centrée sur la question posée sans ajouter de détails non requis."
        )

        custom_prompt_template = custom_prompt_template_general if self.general else custom_prompt_template_personnal
        completion = self.client.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[
                {"role": "system", "content": custom_prompt_template},
                {"role": "user", "content": query}
            ]
        )
        response_message = completion.choices[0].message.content
        return response_message

class IntentionDetector:
    def __init__(self, model_name='all-MiniLM-L6-v2'):
        # Initialize the model
        self.model = SentenceTransformer(model_name)

        # Define the phrases to be used for training the model
        self.phrases_renseigner = [
            "Je voudrais des informations sur le cabinet",
            "Pouvez-vous me renseigner sur vos services?",
            "Quels sont les domaines d'expertise du cabinet?",
            "J'aimerais en savoir plus sur le cabinet",
            "Quelles informations pouvez-vous me donner?",
            "Quelles sont les heures d'ouverture?",
            "Où se situe votre cabinet?",
            "Comment fonctionne votre cabinet?",
            "Quels sont les moyens de paiement acceptés par le cabinet ?",
            "Acceptez-vous les cartes bancaires, ou seulement les chèques ?",
            "Puis-je régler mes consultations par virement bancaire ou PayPal ?",
            "Est-ce que vous acceptez les paiements en ligne ?",
            "Quels types de paiements sont possibles pour une consultation ?",
            "Acceptez-vous les assurances ou les mutuelles pour couvrir les frais ?",
            "Quels sont vos services?",
            "Quel est votre domaine de spécialisation?",
            "Est-ce que le cabinet prend en charge de nouveaux clients?",
            "Quels types de consultation proposez-vous?",
            "Quels sont les honoraires?",
            "Pouvez-vous m'informer sur les tarifs?",
            "Est-ce que le cabinet est ouvert le week-end?",
            "Avez-vous un site web où je peux consulter vos services?",
            "Comment puis-je contacter le cabinet?",
            "Quels sont vos horaires d'ouverture?",
            "Quels sont les moyens de communication avec le cabinet?",
            "Le cabinet propose-t-il des consultations à distance?",
            "Est-ce que vous acceptez la carte vitale?",
            "Quels sont les délais d'attente pour une consultation?",
            "Comment est-ce que je peux obtenir plus d'informations?",
            "Pouvez-vous m'envoyer des détails sur vos services?",
            "Quelle est votre adresse?",
            "Quelles sont les qualifications des praticiens?",
            "Pouvez-vous me donner des précisions sur vos offres?",
            "Est-ce que des consultations d'urgence sont disponibles?",
            "Combien de temps dure une consultation en moyenne?",
            "Avez-vous des services de suivi?",
            "Est-ce que vous faites des consultations à domicile?",
            "Pourriez-vous me fournir des informations complètes concernant votre cabinet et ses services ?",
            "J'aimerais obtenir plus de renseignements sur le cabinet et ses horaires, ainsi que les services proposés.",
            "Quelles sont les spécialités de votre cabinet et comment puis-je en savoir plus ?",
            "Est-ce que vous avez des informations sur les consultations à distance et les tarifs associés ?",
            "Je suis intéressé par des détails concernant vos offres, pourriez-vous me donner plus de précisions ?",
            "Pourriez-vous m'indiquer où je pourrais trouver plus d'informations sur le cabinet ?",
            "Je cherche à en savoir davantage sur vos horaires et services, auriez-vous des précisions à me donner ?",
            "Le cabinet est-il ouvert pendant les jours fériés ?",
            "Est-ce que vous travaillez lors des jours fériés, notamment à Noël ?",
            "Quelles sont vos heures d'ouverture pendant les vacances et jours fériés ?",
            "Est-ce que vous avez un planning spécial pour les jours fériés ?",
            "Le cabinet reste-t-il fermé pendant les jours fériés ou il y a-t-il une permanence ?",
            "Est-ce que vous êtes ouverts lors du Nouvel An ?",
            "Quel est le tarif moyen pour une consultation ?",
            "Pouvez-vous me donner une estimation des tarifs pour une première consultation ?",
            "Est-ce que vous avez des tarifs différents pour les consultations de suivi ?",
            "Quels sont vos tarifs pour les consultations en ligne ?",
            "Y a-t-il des frais supplémentaires pour les consultations urgentes ?",
            "Je voudrais connaître vos tarifs avant de prendre rendez-vous.",
            "Avez-vous une grille tarifaire que je pourrais consulter ?",
            "Quel est votre numéro de téléphone pour prendre rendez-vous ?",
            "Comment puis-je vous contacter par téléphone pour poser des questions ?",
            "Pouvez-vous me communiquer le numéro de contact pour le cabinet ?",
            "Je souhaite joindre votre cabinet par téléphone, quel est le numéro ?",
            "Comment puis-je obtenir un numéro pour vous contacter directement ?",
            "Est-ce que le cabinet a un numéro de contact spécifique pour les urgences ?",
            "J'aimerais discuter avec un praticien, quel est le moyen de vous joindre ?"
            "Comment prendre rendez-vous ?",
            "Dois-je apporter mes anciens examens ?",
            "Est-ce que la mammographie fait mal ?",
            "Quels documents dois-je apporter pour mon examen ?",
            "Je pense être enceinte, puis-je passer mon examen ?"
            "Est-ce que je dois attendre en salle d'attente le compte rendu du radiologue ?"
        ]

        self.phrases_gerer_rdv = [
            "Gestion des rendez-vous",
            "Annulation de rendez-vous",
            "Modification de rendez-vous",
            "Prise de rendez-vous",
            "Plannification",
            "Réservation ",
            "je souhaite gérer mes rendez-vous",
            "Je veux prendre un rendez-vous",
            "Puis-je modifier mon rendez-vous?",
            "J'aimerais annuler mon rendez-vous",
            "Aidez-moi à gérer mes rendez-vous",
            "Comment puis-je organiser un rendez-vous?",
            "Puis-je déplacer mon rendez-vous?",
            "Je souhaite planifier un rendez-vous",
            "Comment réserver un rendez-vous?",
            "Je voudrais confirmer mon rendez-vous",
            "Pouvez-vous m'aider à programmer un rendez-vous?",
            "Est-il possible de reporter mon rendez-vous?",
            "Je voudrais vérifier la date de mon rendez-vous",
            "J'ai besoin de changer l'horaire de mon rendez-vous",
            "Pouvez-vous me rappeler la date de mon rendez-vous?",
            "Est-il possible de prendre rendez-vous en ligne?",
            "Comment puis-je contacter pour un changement de rendez-vous?",
            "Puis-je reprogrammer ma consultation?",
            "Je veux fixer une nouvelle date de rendez-vous",
            "Puis-je planifier plusieurs rendez-vous?",
            "J'ai besoin de décaler mon rendez-vous",
            "Comment puis-je annuler ma consultation?",
            "Avez-vous des créneaux disponibles?",
            "Est-il possible de confirmer mon rendez-vous?",
            "Je souhaite ajuster l'heure de mon rendez-vous",
            "Est-ce que je peux fixer une consultation cette semaine?",
            "Puis-je réserver une consultation pour un proche?",
            "Comment procéder pour réserver un rendez-vous?",
            "Je dois m'inscrire pour une consultation",
            "Est-ce que le cabinet accepte les rendez-vous urgents?",
            "Avez-vous des créneaux libres dans les prochains jours?",
            "Puis-je planifier une visite de suivi?",
            "Quel est le délai pour obtenir un rendez-vous?",
            "Je voudrais vérifier la disponibilité pour un rendez-vous",
            "Je cherche un rendez-vous le plus tôt possible",
            "Est-ce que je peux annuler ou changer un rendez-vous en ligne?",
            "Puis-je organiser un rendez-vous pour plusieurs personnes?",
            "Quel est le processus pour annuler un rendez-vous?",
            "J'aimerais reprogrammer un rendez-vous pour la semaine prochaine, est-ce possible ?",
            "Pouvez-vous m'aider à annuler ma consultation prévue pour demain et la replanifier à une autre date ?",
            "Je souhaite savoir s'il est possible de modifier l'heure de mon rendez-vous, j'ai un conflit d'agenda.",
            "J'aimerais réserver un créneau pour une consultation, quelles sont vos disponibilités cette semaine ?",
            "Est-ce qu'il y a une manière de confirmer ma réservation de rendez-vous en ligne sans passer par l'accueil ?",
            "Est-ce que je peux reporter mon rendez-vous de jeudi à vendredi sans perdre ma place ?",
            "Est-il possible de changer mon rendez-vous, si oui, comment faire ?",
            "J'aimerais vérifier l'heure exacte de ma consultation prévue, est-ce que vous pouvez m'aider ?",
            "Je voudrais annuler mon rendez-vous, mais je ne suis pas sûr de comment procéder, pouvez-vous m'aider ?"
        ]

        # Combine all phrases
        self.all_phrases = self.phrases_renseigner + self.phrases_gerer_rdv

        # Compute embeddings for all phrases
        self.phrase_embeddings = self.model.encode(self.all_phrases, convert_to_numpy=True)

        # Create a FAISS index
        self.dimension = self.phrase_embeddings.shape[1]
        self.index = faiss.IndexFlatL2(self.dimension)
        self.index.add(self.phrase_embeddings)

    def detect_intention(self, reponse):
        if not isinstance(reponse, str) or not reponse.strip():
            return "Option non reconnue"

        # Encode the user's response
        user_embedding = self.model.encode([reponse], convert_to_numpy=True)

        # Search in the FAISS index
        distances, indices = self.index.search(user_embedding, 1)

        # Retrieve the best match
        best_index = indices[0][0]
        best_distance = distances[0][0]
        if best_index < len(self.phrases_renseigner):
            return "renseigner"
        else:
            return "gérer"

class RAG_general:
    def __init__(self , llm_model="gpt-3.5-turbo"):
        if os.path.exists("faiss_index"):
            self.knowledge_base = KnowledgeBase('')
            self.knowledge_base.retriever = FAISS.load_local("faiss_index", self.knowledge_base.embeddings, allow_dangerous_deserialization=True).as_retriever()
        else:
            documents = self.load_files_contents('data')
            self.knowledge_base = KnowledgeBase(documents)
            self.knowledge_base.build_retriever()

        self.assistant = VirtualAssistant(retriever=self.knowledge_base.retriever,general=True,llm_model=llm_model)
    @staticmethod
    def load_files_contents(data_path):
        concatenated_content = ""
        for filename in os.listdir(data_path):
            filepath = os.path.join(data_path, filename)
            if filename.endswith(".pdf"):
                loader = PyPDFLoader(filepath)
                documents = loader.load()
                concatenated_content += "\n".join([doc.page_content for doc in documents]) + "\n"
            elif filename.endswith(".txt"):
                with open(filepath, 'r', encoding='utf-8') as file:
                    concatenated_content += file.read() + "\n"
            elif filename.endswith(".json"):
                with open(filepath, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    concatenated_content += json.dumps(data, indent=2) + "\n"
            elif filename.endswith(".docx"):
                doc = DocxDocument(filepath)
                concatenated_content += "\n".join([paragraph.text for paragraph in doc.paragraphs]) + "\n"
        return concatenated_content


    def answer_query(self, query):
        relevant_docs = self.knowledge_base.retriever.get_relevant_documents(query)
        context = "\n".join([doc.page_content for doc in relevant_docs])
        response_message = self.assistant.answer_query(query, context)
        return response_message
    
class InformationExtractor:
    def __init__(self):
        # Initialisation unique du modèle NER
        logger.info("Initialisation du modèle NER...")
        self.tokenizer = AutoTokenizer.from_pretrained("Jean-Baptiste/camembert-ner-with-dates")
        self.model = AutoModelForTokenClassification.from_pretrained("Jean-Baptiste/camembert-ner-with-dates")
        self.nlp = pipeline('ner', model=self.model, tokenizer=self.tokenizer, aggregation_strategy="simple")
        logger.info("Modèle NER initialisé avec succès.")

    def check_noun(self, msg_2_check):
        logger.debug(f"Vérification du nom : {msg_2_check}")
        def check_str(msg_2_check: str) -> bool:
            return isinstance(msg_2_check, str) and bool(msg_2_check.strip()) and any(ele in msg_2_check for ele in ["a", "e", "i", "o", "u", "y"])

        if not check_str(msg_2_check):
            logger.warning(f"Le message {msg_2_check} n'est pas une chaîne valide.")
            return False

        if not re.match(r"^[a-zA-ZÀ-ÿ' -]+$", msg_2_check):
            logger.warning(f"Le message {msg_2_check} contient des caractères invalides.")
            return False
        return True

    def extraire_nom(self, texte):
        logger.info(f"Extraction du nom à partir du texte : {texte}")
        entities = self.nlp(texte)
        for ent in entities:
            if ent['entity_group'] == "PER":
                if self.check_noun(ent['word'].lower()):
                    logger.info(f"Nom extrait : {ent['word'].upper()}")
                    return ent['word'].upper()
        logger.warning("Aucun nom n'a été extrait.")
        return None

    def extraire_prenom(self, texte):
        logger.info(f"Extraction du prénom à partir du texte : {texte}")
        entities = self.nlp(texte)
        for ent in entities:
            if ent['entity_group'] == "PER":
                if self.check_noun(ent['word']):
                    logger.info(f"Prénom extrait : {ent['word']}")
                    return ent['word'].upper()
        logger.warning("Aucun prénom n'a été extrait.")
        return None

    def extraire_date_naissance(self, texte):
        logger.info(f"Extraction de la date de naissance à partir du texte : {texte}")
        entities = self.nlp(texte)
        for ent in entities:
            if ent['entity_group'] == "DATE":
                date_str = ent['word']
                date_obj = dateparser.parse(date_str)
                if date_obj:
                    formatted_date = date_obj.strftime("%Y-%m-%d")
                    logger.info(f"Date de naissance extraites : {formatted_date}")
                    return formatted_date
                else:
                    logger.warning(f"Date non valide extraites : {date_str}")
                    return date_str
        logger.warning("Aucune date de naissance n'a été extraite.")
        return None

    def extraire_adresse(self, texte):
        logger.info(f"Extraction de l'adresse à partir du texte : {texte}")
        # Extraction du numéro de rue
        numero_rue = re.search(r'\b\d+\b', texte)
        adr = f"{numero_rue.group()} " if numero_rue else ""
        adr=''
        # Extraction des entités pertinentes
        entities = self.nlp(texte)
        for ent in entities:
            if ent['entity_group'] in {"LOC", "PER"}:
                adr += ent['word'] + ' '

        adr = adr.strip()
        if adr:
            logger.info(f"Adresse extraite : {adr}")
        else:
            logger.warning("Aucune adresse n'a été extraite.")
        return adr

    def extraire_numero_telephone(self, texte):
        logger.info(f"Extraction du numéro de téléphone à partir du texte : {texte}")
        # Normalisation du texte en supprimant les espaces, tirets, et points
        phone_number = texte.replace(" ", "").replace("-", "").replace(".", "")

        # Premier regex : validation des numéros compactés
        phone_regex = r"^(\+?\d{1,3})?(\d{9,10})$"
        numero_telephone = re.search(phone_regex, phone_number)

        if numero_telephone:
            logger.info(f"Numéro de téléphone extrait : {numero_telephone.group()}")
            return numero_telephone.group()

        # Deuxième regex : validation des formats avec séparateurs (espaces, tirets)
        numero_telephone = re.search(r"(\+?\d{1,3}[\s-]?)?(\(?\d{1,4}\)?[\s-]?)?(\d{2}[\s-]?){4}\d{2}", phone_number)

        if numero_telephone:
            logger.info(f"Numéro de téléphone extrait avec séparateurs : {numero_telephone.group()}")
            return numero_telephone.group()

        logger.warning("Aucun numéro de téléphone valide n'a été extrait.")
        return None

    def extraire_code_postal(self, texte):
        logger.info(f"Extraction du code postal à partir du texte : {texte}")
        code_postal = re.search(r"\b\d{5}\b", texte)
        if code_postal:
            logger.info(f"Code postal extrait : {code_postal.group()}")
            return code_postal.group()
        else:
            logger.warning("Aucun code postal valide n'a été extrait.")
        return None

    def extraire_adresse_mail(self, texte):
        logger.info(f"Extraction de l'adresse email à partir du texte : {texte}")
        texte = re.sub(r'\s*arobase\s*', '@', texte, flags=re.IGNORECASE)
        adresse_mail = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", texte)

        if adresse_mail:
            logger.info(f"Adresse email extraite : {adresse_mail[0].strip()}")
            return adresse_mail[0].strip()
        else:
            logger.warning("Aucune adresse email valide n'a été extraite.")
        return None

class RendezVousManager:
    def __init__(self, nom, prenom, date_naissance, email):
        self.tokenizer = AutoTokenizer.from_pretrained("Jean-Baptiste/camembert-ner-with-dates")
        self.model = AutoModelForTokenClassification.from_pretrained("Jean-Baptiste/camembert-ner-with-dates")
        self.nlp = pipeline('ner', model=self.model, tokenizer=self.tokenizer, aggregation_strategy="simple")
        self.nom = nom
        self.prenom = prenom
        self.date_naissance = date_naissance
        self.email = email
        self.rdv_prevus = get_rdv_prevus(self.nom, self.prenom, self.date_naissance, self.email)
        self.prendre_rdv = self.PrendreRdv(self)
        self.modifier_rdv = self.ModifierRdv(self)

        logger.info(f'RendezVousManager initialized for {self.nom} {self.prenom}.')

    def annuler_rdv(self, choix_patient='', rdv_ID=None):
        logger.info('Attempting to cancel an appointment.')
        if rdv_ID:
            logger.info(f'Removing appointment with ID {rdv_ID}.')
            return remove_rdv(self.nom, self.prenom, self.date_naissance, self.email, rdv_ID)
        else:
            date_annuler = get_creneau(self.nlp, choix_patient)
            if date_annuler:
                logger.info(f'Cancelling appointment for the date {date_annuler}.')
                if self.rdv_prevus:
                    for rdv_id, details in self.rdv_prevus.items():
                        if details['date'] == date_annuler:
                            logger.info(f'Found matching appointment. Removing appointment ID {rdv_id}.')
                            return remove_rdv(self.nom, self.prenom, self.date_naissance, self.email, rdv_id)
                else:
                    # Si self.rdv_prevus est vide ou aucun rendez-vous ne correspond à la date
                    logger.warning(f'No appointments scheduled for {date_annuler}. Cannot cancel.')  # Log plus spécifique
            else:
                  logger.warning('No valid date provided for cancellation.')
            return False

    class ModifierRdv:
        def __init__(self, parent):
            self.parent = parent

        def get_creneaux_dispo_pour_modification(self, choix_patient):
            logger.info('Attempting to get available slots for modification.')
            date_modifier = get_creneau(self.parent.nlp, choix_patient)
            if date_modifier:
                if self.parent.rdv_prevus:
                    for rdv_id, details in self.parent.rdv_prevus.items():
                        if details['date'] == date_modifier:
                            logger.info(f'Found appointment to modify with ID {rdv_id}.')
                            creneaux_disponibles = get_creneaux_dispo(
                                details['type'], details['code'], details['date']
                            )
                            return creneaux_disponibles, rdv_id
            logger.warning('No available slots for modification found.')
            return None, None

        def modifier_rdv(self, choix_patient, creneaux_dispo, rdv_ID):
            logger.info('Attempting to modify appointment.')
            date_modifier = get_creneau(self.parent.nlp, choix_patient)
            if date_modifier:
                if creneaux_dispo:
                    for cr_ID, details in creneaux_dispo.items():
                        if date_modifier == (datetime.strptime(details['date'], '%Y-%m-%dT%H:%M:%S').replace(hour=int(details['heureDebut'][:2]), minute=int(details['heureDebut'][3:5]))).strftime('%Y-%m-%dT%H:%M:%S'):
                            logger.info(f'Found matching slot for modification. Updating appointment ID {rdv_ID} with slot ID {cr_ID}.')
                            return update_rdv(self.parent.nom, self.parent.prenom, self.parent.date_naissance, self.parent.email, rdv_ID, cr_ID)
            logger.warning('No matching slot found for modification.')
            return False

def pipeline_lyrae_talk():
    analyzer = AnalyseurConversation()

    # Recueil du consentement
    consentement = analyzer.recueil_consentement(input("Cette communication est susceptible d'être enregistrée. Acceptez-vous de parler avec moi ? Merci de répondre par Oui ou Non."))
    if consentement['status'] == 'Consentement valide':
        print("Merci pour votre consentement. Nous pouvons continuer.")
        detector = IntentionDetector()

        # Détection de l'intention
        intention = detector.detect_intention(input("Que souhaitez-vous faire ? Si vous voulez obtenir des informations sur le cabinet, dites renseigner. Si vous souhaitez gérer vos rendez-vous, dites gérer mes rendez-vous. ?"))
        if intention == 'renseigner':
            rag_system1 = RAG_general()
            while not analyzer.quitter_conversation(input("Voulez-vous quitter la conversation ? (par exemple, 'J'ai fini ma discussion') : ")):
                query = input("Quelle est votre question ? : ")
                response = rag_system1.answer_query(query)
                print(f"Réponse : {response}")
        else:
            patient_ident = InformationExtractor()

            # Extraction des informations personnelles
            nom = patient_ident.extraire_nom(input("Quel est votre nom de famille ? : "))
            prenom = patient_ident.extraire_prenom(input("Quel est votre prénom ? : "))
            date_naissance = patient_ident.extraire_date_naissance(input("Quelle est votre date de naissance ? : "))
            email = patient_ident.extraire_adresse_mail(input("Quelle est votre adresse e-mail ? : "))
            rdv_prevus= get_rdv_prevus(nom, prenom, date_naissance, email)
            if get_user(nom, prenom, date_naissance, email):
                print("Patient connu")
                manager = RendezVousManager(nom, prenom, date_naissance, email)

                # Demande du choix d'action
                choix = gestion_rdv(input("Que souhaitez-vous faire ? (prendre, modifier et annuler un rendez-vous ou consulter) : ").lower())
                if choix == 'annuler':
                    # Annulation du rendez-vous
                    success = manager.annuler_rdv(input("Quel rendez-vous souhaitez-vous annuler ?"))
                    if success:
                        print("Annulation effectuée")
                    else:
                        print("Annulation échouée")
                elif choix == 'modifier':
                    # Modification du rendez-vous
                    rdv_description = input("Quel rendez-vous souhaitez-vous modifier ? : ")
                    creneaux, rdv_id = manager.get_creneaux_dispo_pour_modification(rdv_description)
                    if creneaux:
                        print(f"Vous souhaitez modifier un rendez-vous. Voici les créneaux disponibles : {creneaux}")
                        nouveau_creneau = input("Quel créneau souhaitez-vous choisir ? : ")
                        if manager.modifier_rdv(nouveau_creneau, creneaux, rdv_id):
                            print("Modification effectuée avec succès")
                        else:
                            print("Modification échouée")
                    else:
                        print("Aucun créneau disponible pour la modification")
                elif choix == 'prendre':
                    # Prise de rendez-vous
                    date_debut = manager.prendre_rdv.process1_prendre_rdv(input("À partir de quelle date souhaitez-vous prendre un rendez-vous ? : "))
                    type_examen =manager.prendre_rdv.process2_prendre_rdv(input("Quel est le type de l'examen pour lequel vous voulez prendre rendez-vous ? IRM, Scanner, Radiologie ou Échographie ?"))
                    code_examen =type_examen, manager.prendre_rdv.process3_prendre_rdv(type_examen ,input("Pour quelle partie anatomique ? (par exemple, 'bassin', 'genou', etc.) : "))
                    creneaux_dispo = get_creneaux_dispo(type_examen,code_examen, date_debut)
                    if creneaux_dispo:
                        print(f"Vous souhaitez prendre un rendez-vous. Voici les créneaux disponibles :{creneaux_dispo}")
                        choix_creneau = input("Quel créneau horaire souhaitez-vous réserver ? Indiquez le jour, le mois, l'année, l'heure et les minutes.")
                        if manager.prendre_rdv.process4_prendre_rdv(creneaux_dispo, choix_creneau):
                            print("Prise de rendez-vous effectuée")
                        else:
                            print("Prise de rendez-vous échouée")
                    else:
                        print("Désolé , notre planning est complet pour ce type d'examen. Vous pourrez prendre RDV à la publication du prochain planning")
                elif choix == 'consulter':
                  rag_system2 = RAG_personnal(context=f"Les rendez vous que vous avez réservés sont : {str(get_rdv_prevus(nom,prenom,date_naissance,email))}")
                  while not analyzer.quitter_conversation(input("Voulez-vous quitter la conversation ? (par exemple, 'J'ai fini ma discussion') : ")):
                      query = input("Quelle est votre question ? : ")
                      response = rag_system2.answer_query(query)
                      print(f"Réponse : {response}")
            else:
                print("Patient inconnu")
    else:
        print("Consentement non valide ou refusé")

class RAG_personnal:
    def __init__(self ,context, llm_model="gpt-3.5-turbo"):
            self.knowledge_base = KnowledgeBase(context)
            self.knowledge_base.build_retriever(save=False)
            self.assistant = VirtualAssistant(retriever=self.knowledge_base.retriever,general=False,llm_model=llm_model)
    def answer_query(self, query):
        relevant_docs = self.knowledge_base.retriever.get_relevant_documents(query)
        context = "\n".join([doc.page_content for doc in relevant_docs])
        response_message = self.assistant.answer_query(query, context)
        return response_message

def get_rdv_prevus (nom, prenom, date_naissance,email):
    return {
        "rdv1": {"date": '2025-01-10T15:30:00', "type": "CT",'code':f'N01CTABD'},
        "rdv2": {"date": '2025-03-03T19:20:00',  "type": "CT",'code':f'N01CTABD'},
        "rdv3": {"date": '2025-06-10T19:10:00',  "type": "CT",'code':f'N01CTABD'}
    }

def get_user(nom, prenom,date_naissance ,email):
      url = "https://ai2xplore.azurewebsites.net/api/getUser"
      payload = json.dumps({
        "email": email,
        "firstName": prenom,
        "lastName": nom,
        "birthDate": date_naissance
      })
      headers = {
        'Content-Type': 'application/json'
      }
      response = requests.request("POST", url, headers=headers, data=payload)
      if response.status_code == 200 :
        return response.json()
      else:
          return False
      
def gestion_rdv(question):
    try:
        # Normalisation de la question
        logger.info(f'Normalizing question: "{question}"')
        question = unicodedata.normalize('NFD', question).encode('ascii', 'ignore').decode('utf-8').lower()
        logger.debug(f'Normalized question: "{question}"')

        # Recherche des verbes et noms relatifs à "annuler"
        if re.search(r"\b(annuler|supprimer|effacer|abandonner|annulation|suppression|abondon|effacement)\b", question):
            logger.info(f'Found "annuler" related keywords in question: "{question}"')
            return "annuler"

        # Recherche des verbes et noms relatifs à "modifier"
        elif re.search(r"\b(modifier|changer|deplacer|reorganiser|reporter|ajuster|corriger|amender|modification|changement|ajustement|reprogrammer|deplacement|correction)\b", question):
            logger.info(f'Found "modifier" related keywords in question: "{question}"')
            return "modifier"

        # Recherche des verbes et noms relatifs à "prendre un rendez-vous"
        elif re.search(r"\b(prendre|reserver|planifier|organiser|reservation|créneau|programmer|planification|prise|programmation|organisation)\b", question):
            logger.info(f'Found "prendre" related keywords in question: "{question}"')
            return "prendre"

        # Si aucune correspondance n'est trouvée
        else:
            logger.info(f'No matching keywords found. Returning "consulter" for question: "{question}"')
            return "consulter"

    except Exception as e:
        logger.error(f"Error while processing the question: {e}")
        return "invalide"
    
def get_creneaux_dispo (type_examen, code_examen, date_debut):
      url = "https://ai2xplore.azurewebsites.net/api/getCreneaux"
      payload = json.dumps({
        "typeExamen" :type_examen ,
        "codeExamen":code_examen,
        "dateDebut" :date_debut
      })
      headers = {
        'Content-Type': 'application/json'
      }
      response = requests.request("POST", url, headers=headers, data=payload)
      if response.status_code == 200 :
        return response.json()
      else:
          return False
      
class KnowledgeBase:
    def __init__(self, text_data):
        texts=self.text_split(text_data)
        self.documents = self.convert_texts_to_documents(texts)
        self.embeddings = OpenAIEmbeddings(model="text-embedding-ada-002",api_key= os.environ["OPENAI_API_KEY"])
        self.retriever = None

    @staticmethod
    def text_split(extracted_data):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size =200, chunk_overlap =50,separators=["\n"])
        text_chunks = text_splitter.split_text(extracted_data)
        return text_chunks

    def build_retriever(self,save=True):
        store = FAISS.from_documents(self.documents, self.embeddings)
        if save:
               store.save_local("faiss_index")
        self.retriever =store.as_retriever()

    @staticmethod
    def convert_texts_to_documents(texts):
        documents = [Document(page_content=text) for text in texts]
        return documents

def remove_rdv(nom, prenom, date_naissance,email, rdv_id):
      url = "https://ai2xplore.azurewebsites.net/api/deleteRDV"
      payload = json.dumps({
        "rdvId":rdv_id,
        "email": email,
        "firstName": prenom,
        "lastName": nom,
        "birthDate": date_naissance,
        "externalUserNumber": "00000"

      })
      headers = {
        'Content-Type': 'application/json'
      }
      response = requests.request("DELETE", url, headers=headers, data=payload)
      if response.status_code == 200 :
               return response.json()
      else :
            return False

def get_creneau(nlp ,choix_patient):
    logger.info('Attempting to extract appointment slot from the input text.')

    def convert_french_numbers_to_digits(text):
        logger.debug(f'Converting French numbers in the text: "{text}"')
        french_number_mapping = {
            "premier": "1", "un": "1", "deux": "2", "trois": "3", "quatre": "4", "cinq": "5",
            "six": "6", "sept": "7", "huit": "8", "neuf": "9", "dix": "10",
            "onze": "11", "douze": "12", "treize": "13", "quatorze": "14", "quinze": "15",
            "seize": "16", "dix-sept": "17", "dix-huit": "18", "dix-neuf": "19",
            "vingt": "20", "vingt et un": "21", "vingt-et-un": "21",
            "vingt-deux": "22", "vingt trois": "23", "vingt-trois": "23",
            "trente": "30", "trente et un": "31", "trente-et-un": "31", 'minuit':'24h',
            'midi': '12','deux heures': '2h', 'trois heures': '3h', 'quatre heures': '4h',
          'cinq heures': '5h', 'six heures': '6h', 'sept heures': '7h', 'huit heures': '8h',
          'neuf heures': '9h', 'dix heures': '10h', 'onze heures': '11h', 'douze heures': '12h',
          'treize heures': '13h', 'quatorze heures': '14h', 'quinze heures': '15h',
          'seize heures': '16h', 'dix-sept heures': '17h', 'dix-huit heures': '18h',
          'dix-neuf heures': '19h', 'vingt heures': '20h', 'vingt et une heures': '21h',
          'vingt-deux heures': '22h', 'vingt-trois heures': '23h',
        }
        # Remplacer "h" par "h00" si nécessaire
        if re.search(r'(\d{1,2})h(?!\d)', text):
            text = re.sub(r'(\d{1,2})h', r'\1h00', text)
            logger.debug(f'Updated time format with "h00": {text}')

        # Appliquer le remplacement pour les numéros en français
        pattern = re.compile(r'\b(' + '|'.join(re.escape(word) for word in french_number_mapping.keys()) + r')\b', re.IGNORECASE)
        def replace_match(match):
            word = match.group(0).lower()
            replacement = french_number_mapping.get(word, word)
            logger.debug(f'Replacing "{word}" with "{replacement}"')
            return replacement

        return pattern.sub(replace_match, text)

    # Traitement du choix du patient
    logger.info(f'Processing input: "{choix_patient}"')
    entities = nlp(choix_patient)
    logger.debug(f'Entities detected: {entities}')

    creneau_choisi = ''
    for ent in entities:
        if ent['entity_group'] == "DATE":
            creneau_choisi = convert_french_numbers_to_digits(str(ent['word']))
            logger.info(f'Extracted date: "{creneau_choisi}"')

    # Parser la date
    logger.debug(f'Parsing the extracted date: "{creneau_choisi}"')
    date_obj = dateparser.parse(creneau_choisi, languages=['fr'])
    if date_obj:
        formatted_date = date_obj.strftime("%Y-%m-%dT%H:%M:%S")
        logger.info(f'Formatted date: {formatted_date}')
        return formatted_date
    else:
        logger.warning(f'Failed to parse date from input: "{creneau_choisi}"')
        return None

def update_rdv(nom, prenom, date_naissance,email,rdv_id ,cr_id) :
    return True
