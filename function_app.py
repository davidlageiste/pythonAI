# function_app.py

import azure.functions as func
import logging
from langchain.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader, JSONLoader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from openai import OpenAI
from docx import Document as DocxDocument
import os
import json
from azure.storage.blob import BlobServiceClient

class KnowledgeBase:
    def __init__(self, text_data):
        texts = self.text_split(text_data)
        self.documents = self.convert_texts_to_documents(texts)
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-ada-002",
            api_key=os.environ["OPENAI_API_KEY"]
        )
        self.retriever = None

    @staticmethod
    def text_split(extracted_data):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=200, 
            chunk_overlap=50,
            separators=["\n"]
        )
        text_chunks = text_splitter.split_text(extracted_data)
        return text_chunks

    def build_retriever(self):
        store = FAISS.from_documents(self.documents, self.embeddings)
        # Save to Azure Blob Storage instead of local filesystem
        blob_service = BlobServiceClient.from_connection_string(
            os.environ["AzureWebJobsStorage"]
        )
        container_client = blob_service.get_container_client("faiss-indexes")
        
        # Ensure container exists
        if not container_client.exists():
            container_client.create_container()
        
        # Save FAISS index to temporary file then upload
        store.save_local("temp_faiss_index")
        with open("temp_faiss_index/index.faiss", "rb") as data:
            blob_client = container_client.get_blob_client("index.faiss")
            blob_client.upload_blob(data, overwrite=True)
        
        with open("temp_faiss_index/index.pkl", "rb") as data:
            blob_client = container_client.get_blob_client("index.pkl")
            blob_client.upload_blob(data, overwrite=True)
        
        self.retriever = store.as_retriever()

    @staticmethod
    def convert_texts_to_documents(texts):
        documents = [Document(page_content=text) for text in texts]
        return documents

class VirtualAssistant:
    def __init__(self, retriever, general=True, llm_model="gpt-3.5-turbo"):
        self.retriever = retriever
        self.client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
        self.general = general

    def answer_query(self, query, context):
        custom_prompt_template = (
            f"Vous êtes une assistante virtuelle spécialisée pour un cabinet de radiologie médicale.\n"
            f"Voici des informations pertinentes extraites de notre base de connaissances :\n{context}\n\n"
            f"Question du patient : {query}\n\n"
            f"Répondez uniquement à la question posée, de manière claire et concise."
        )

        completion = self.client.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[
                {"role": "system", "content": custom_prompt_template},
                {"role": "user", "content": query}
            ]
        )
        return completion.choices[0].message.content

class RAG_Azure:
    def __init__(self, llm_model="gpt-3.5-turbo"):
        # Try to load from Azure Blob Storage
        blob_service = BlobServiceClient.from_connection_string(
            os.environ["AzureWebJobsStorage"]
        )
        container_client = blob_service.get_container_client("faiss-indexes")
        
        try:
            # Download FAISS index from blob storage
            blob_client = container_client.get_blob_client("index.faiss")
            with open("/tmp/index.faiss", "wb") as f:
                f.write(blob_client.download_blob().readall())
            
            blob_client = container_client.get_blob_client("index.pkl")
            with open("/tmp/index.pkl", "wb") as f:
                f.write(blob_client.download_blob().readall())
            
            self.knowledge_base = KnowledgeBase('')
            self.knowledge_base.retriever = FAISS.load_local(
                "/tmp", 
                self.knowledge_base.embeddings
            ).as_retriever()
            
        except Exception as e:
            logging.info(f"No existing index found, creating new one: {str(e)}")
            # Load documents from blob storage
            documents = self.load_documents_from_blob()
            self.knowledge_base = KnowledgeBase(documents)
            self.knowledge_base.build_retriever()

        self.assistant = VirtualAssistant(
            retriever=self.knowledge_base.retriever,
            general=True,
            llm_model=llm_model
        )

    @staticmethod
    def load_documents_from_blob():
        blob_service = BlobServiceClient.from_connection_string(
            os.environ["AzureWebJobsStorage"]
        )
        container_client = blob_service.get_container_client("documents")
        
        concatenated_content = ""
        
        # List all blobs in the container
        blobs = container_client.list_blobs()
        for blob in blobs:
            blob_client = container_client.get_blob_client(blob.name)
            content = blob_client.download_blob().readall()
            
            if blob.name.endswith('.pdf'):
                # Save temporarily and load with PyPDFLoader
                with open("/tmp/temp.pdf", "wb") as f:
                    f.write(content)
                loader = PyPDFLoader("/tmp/temp.pdf")
                documents = loader.load()
                concatenated_content += "\n".join([doc.page_content for doc in documents])
            
            elif blob.name.endswith('.txt'):
                concatenated_content += content.decode('utf-8')
            
            elif blob.name.endswith('.json'):
                data = json.loads(content)
                concatenated_content += json.dumps(data, indent=2)
            
            elif blob.name.endswith('.docx'):
                # Save temporarily and load with DocxDocument
                with open("/tmp/temp.docx", "wb") as f:
                    f.write(content)
                doc = DocxDocument("/tmp/temp.docx")
                concatenated_content += "\n".join([p.text for p in doc.paragraphs])
        
        return concatenated_content

    def process_query(self, query):
        retrieved_docs = self.knowledge_base.retriever.get_relevant_documents(query)
        context = "\n".join([doc.page_content for doc in retrieved_docs])
        return self.assistant.answer_query(query, context)


# Azure Function
app = func.FunctionApp()

@app.route(route="rag_query")
def rag_query(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    try:
        req_body = req.get_json()
        query = req_body.get('text')

        if not query:
            return func.HttpResponse(
                json.dumps({"error": "No query provided in request body"}),
                mimetype="application/json",
                status_code=400
            )

        rag_system = RAG_Azure()
        result = rag_system.process_query(query)

        return func.HttpResponse(
            json.dumps({"response": result}),
            mimetype="application/json"
        )

    except Exception as e:
        logging.error(f"Error processing request: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            status_code=500
        )